#!/usr/bin/env python3
"""Finalize GridVision copyright Word document (copy only, never touch source)."""
from __future__ import annotations

import re
import shutil
import sys
from pathlib import Path

SRC = Path("GridVision_软件著作权代码最终提交版_Consolas版.docx")
DST = Path("GridVision_软件著作权代码最终版.docx")

FORBIDDEN_PHRASES = (
    "源文件:",
    "源文件：",
    "来源:",
    "来源：",
    "行号:",
    "行号：",
    "本文件为摘录",
    "完整依赖见原工程",
    "GridVision 毕设答辩核心代码摘录",
    "Clean版",
    "净化版",
    "净化统计",
    "调试统计",
    "代码统计",
    "Clean",
    "摘录",
    "audit",
    "debug",
)

AUDIT_CALLS = (
    "logInspectionPointCoordAudit",
    "logInspectionRouteBindingSummary",
    "logRoutePointAudit",
    "debugSummary",
    "auditSummary",
)

AUDIT_CALL_PATTERN = re.compile(
    r"^\s*(_log_[a-zA-Z0-9_]+|_point_flow_log_dashboard|"
    r"logInspectionPointCoordAudit|logInspectionRouteBindingSummary|"
    r"logRoutePointAudit|debugSummary|auditSummary)\s*\("
)

STATS = {
    "console_log": 0,
    "console_warn": 0,
    "console_error": 0,
    "print": 0,
    "audit_calls": 0,
}


def set_run_font(run, name: str, size_pt: float | None = None, bold: bool = False):
    from docx.oxml.ns import qn

    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        from docx.shared import Pt

        run.font.size = Pt(size_pt)
    run.bold = bold


def add_centered_title(doc, lines: list[str], size: float = 22):
    for text in lines:
        p = doc.add_paragraph()
        p.alignment = 1  # CENTER
        run = p.add_run(text)
        set_run_font(run, "黑体", size, bold=True)


def add_heading_black(doc, text: str, size: float = 16):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "黑体", size, bold=True)


def add_label_line(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "黑体", 11, bold=True)


def add_desc_line(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, "宋体", 11)


def add_code_block(doc, code: str):
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt

    for line in code.split("\n"):
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_run_font(run, "Consolas", 9)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)


def contains_forbidden(text: str) -> bool:
    low = text.lower()
    for phrase in FORBIDDEN_PHRASES:
        if phrase.lower() in low:
            return True
    return False


def remove_call_statement(text: str, start: int) -> tuple[int, str | None]:
    """Remove a call starting at `start`; return (new_index, console_kind|print|audit)."""
    n = len(text)
    m = re.match(r"console\.(log|warn|error)\s*\(", text[start:])
    kind = None
    if m:
        kind = m.group(1)
        j = start + m.end()
    else:
        m2 = re.match(r"print\s*\(", text[start:])
        if m2:
            kind = "print"
            j = start + m2.end()
        else:
            for name in AUDIT_CALLS:
                m3 = re.match(rf"{re.escape(name)}\s*\(", text[start:])
                if m3:
                    kind = "audit"
                    j = start + m3.end()
                    break
            else:
                return start, None
    depth = 1
    in_str = None
    esc = False
    while j < n and depth > 0:
        c = text[j]
        if in_str:
            if esc:
                esc = False
            elif c == "\\":
                esc = True
            elif c == in_str:
                in_str = None
        elif c in ("'", '"'):
            in_str = c
        elif c == "(":
            depth += 1
        elif c == ")":
            depth -= 1
        j += 1
    while j < n and text[j] in " \t":
        j += 1
    if j < n and text[j] == ";":
        j += 1
    return j, kind


def remove_console_calls(text: str) -> str:
    out = []
    i = 0
    n = len(text)
    while i < n:
        m = re.match(r"console\.(log|warn|error)\s*\(", text[i:])
        if not m:
            out.append(text[i])
            i += 1
            continue
        kind = m.group(1)
        STATS[f"console_{kind}"] += 1
        j = i + m.end()
        depth = 1
        in_str = None
        esc = False
        while j < n and depth > 0:
            c = text[j]
            if in_str:
                if esc:
                    esc = False
                elif c == "\\":
                    esc = True
                elif c == in_str:
                    in_str = None
            elif c in ("'", '"', "`"):
                in_str = c
            elif c == "(":
                depth += 1
            elif c == ")":
                depth -= 1
            j += 1
        while j < n and text[j] in " \t":
            j += 1
        if j < n and text[j] == ";":
            j += 1
        i = j
    return "".join(out)


def remove_print_calls(text: str) -> str:
    out = []
    i = 0
    n = len(text)
    while i < n:
        m = re.match(r"print\s*\(", text[i:])
        if not m:
            out.append(text[i])
            i += 1
            continue
        STATS["print"] += 1
        j = i + m.end()
        depth = 1
        in_str = None
        esc = False
        while j < n and depth > 0:
            c = text[j]
            if in_str:
                if esc:
                    esc = False
                elif c == "\\":
                    esc = True
                elif c == in_str:
                    in_str = None
            elif c in ("'", '"'):
                in_str = c
            elif c == "(":
                depth += 1
            elif c == ")":
                depth -= 1
            j += 1
        while j < n and text[j] in " \t":
            j += 1
        if j < n and text[j] == ";":
            j += 1
        if j < n and text[j] == "\n":
            j += 1
        i = j
    return "".join(out)


def post_clean_code(text: str) -> str:
    text = re.sub(
        r"\s*const segSummary = \[\];\s*let lastKey = null;\s*"
        r"sequence\.forEach\(\(f\) => \{[\s\S]*?\}\);\s*"
        r"const first = sequence\[0\];\s*",
        "\n",
        text,
        count=1,
    )
    text = re.sub(r"\s*const _badSnappedLogged = new Set\(\);\s*\n", "\n", text)
    text = re.sub(r"\s*_badSnappedLogged\.clear\(\);\s*\n", "\n", text)
    text = re.sub(
        r"  const pid = String\(p\.point_id \|\| p\.id \|\| [^}]+\}\s*",
        "",
        text,
        count=1,
    )
    text = re.sub(r"\s*printPlaybackCheck\(\);\s*\n", "\n", text)
    text = re.sub(
        r"function printPlaybackCheck\s*\(\)\s*\{[\s\S]*?\n\}\s*",
        "",
        text,
        count=1,
    )
    text = re.sub(
        r"^\s*`?\[mission-route\][^\n]*\n?",
        "",
        text,
        flags=re.MULTILINE,
    )
    text = re.sub(
        r"^\s*_segments=\$\{[^\n]*\n?",
        "",
        text,
        flags=re.MULTILINE,
    )
    text = re.sub(r"\}\s*else\s*\{\s*\}", "}", text)
    text = re.sub(r"^\s*else\s*\{\s*\}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*if\s*\([^)]*\)\s*\{\s*\}\s*$", "", text, flags=re.MULTILINE)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text


def clean_code(code: str) -> str:
    code = remove_console_calls(code)
    code = remove_print_calls(code)
    lines_out: list[str] = []
    skip_until = 0
    i = 0
    raw_lines = code.split("\n")

    while i < len(raw_lines):
        line = raw_lines[i]
        stripped = line.strip()

        if contains_forbidden(line):
            i += 1
            continue

        if any(x in line for x in ("来源：", "来源:", "源文件：", "源文件:", "行号：", "行号:")):
            i += 1
            continue

        if re.search(r"raw=\$\{rawCount\}|segment=\$\{ev\.segment_id\}", line):
            i += 1
            continue

        if AUDIT_CALL_PATTERN.match(stripped):
            STATS["audit_calls"] += 1
            i += 1
            continue

        if re.search(
            rf"function\s+({'|'.join(AUDIT_CALLS)})\s*\(",
            stripped,
        ):
            depth = 0
            j = i
            while j < len(raw_lines):
                depth += raw_lines[j].count("{") - raw_lines[j].count("}")
                j += 1
                if depth <= 0 and j > i:
                    break
            STATS["audit_calls"] += 1
            i = j
            continue

        if re.search(r"console\.(log|warn|error)\s*\(", line):
            i += 1
            continue

        if re.match(r"^\s*print\s*\(", line):
            i += 1
            continue

        lines_out.append(line)
        i += 1

    text = "\n".join(lines_out)
    text = post_clean_code(text)
    return text.rstrip() + "\n"


def parse_modules(doc) -> tuple[list[dict], list[dict]]:
    frontend: list[dict] = []
    backend: list[dict] = []
    section = None
    pending: dict | None = None

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        if t in ("前端代码",):
            section = "frontend"
            pending = None
            continue
        if t in ("后端代码",):
            section = "backend"
            pending = None
            continue
        if re.match(r"^文件(名称)?[：:]", t):
            name = re.sub(r"^文件(名称)?[：:]", "", t).strip()
            name = re.sub(r"^(frontend|backend)/", "", name)
            pending = {"name": name, "desc": "", "code": ""}
            continue
        if re.match(r"^来源[：:]", t):
            continue
        if re.match(r"^功能(说明)?[：:]", t):
            if pending is not None:
                pending["desc"] = re.sub(r"^功能(说明)?[：:]", "", t).strip()
            continue
        if pending is not None and (
            len(t) > 80
            or t.startswith(("const ", "function ", "def ", "class ", "import ", "from ", "async ", "@"))
        ):
            pending["code"] = t
            if section == "frontend":
                frontend.append(pending)
            elif section == "backend":
                backend.append(pending)
            pending = None
    return frontend, backend


def build_document(frontend: list[dict], backend: list[dict]):
    from docx import Document

    doc = Document()
    add_centered_title(doc, ["GridVision 无人机巡检系统", "", "软件著作权源程序文档"], size=24)
    doc.add_page_break()

    add_heading_black(doc, "前端代码", size=18)
    for idx, mod in enumerate(frontend):
        if idx:
            doc.add_paragraph()
        add_label_line(doc, f"文件名称：{mod['name']}")
        doc.add_paragraph()
        add_label_line(doc, "功能说明：")
        add_desc_line(doc, mod["desc"])
        doc.add_paragraph()
        add_code_block(doc, mod["code"])

    doc.add_page_break()
    add_heading_black(doc, "后端代码", size=18)
    for idx, mod in enumerate(backend):
        if idx:
            doc.add_paragraph()
        add_label_line(doc, f"文件名称：{mod['name']}")
        doc.add_paragraph()
        add_label_line(doc, "功能说明：")
        add_desc_line(doc, mod["desc"])
        doc.add_paragraph()
        add_code_block(doc, mod["code"])

    doc.save(DST)


def final_check(text: str) -> list[str]:
    issues = []
    if "console.log" in text:
        issues.append("console.log")
    if "console.warn" in text:
        issues.append("console.warn")
    if "console.error" in text:
        issues.append("console.error")
    if "print(" in text:
        issues.append("print(")
    for phrase in ("Clean", "摘录", "源文件", "行号", "来源：", "来源:", "净化", "调试统计", "代码统计"):
        if phrase in text:
            issues.append(phrase)
    if AUDIT_CALL_PATTERN.search(text):
        issues.append("audit_call")
    for name in AUDIT_CALLS:
        if name in text:
            issues.append(name)
    return issues


def count_pages() -> int | None:
    try:
        import win32com.client  # type: ignore

        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(DST))
        pages = int(doc.ComputeStatistics(2))
        doc.Close(False)
        word.Quit()
        return pages
    except Exception:
        return None


def main() -> int:
    if not SRC.exists():
        print(f"Source not found: {SRC}", file=sys.stderr)
        return 1

    shutil.copy2(SRC, DST)

    from docx import Document

    src_doc = Document(str(SRC))
    frontend, backend = parse_modules(src_doc)

    for group in (frontend, backend):
        for mod in group:
            mod["code"] = clean_code(mod["code"])

    build_document(frontend, backend)

    out_text = "\n".join(par.text for par in Document(str(DST)).paragraphs)
    issues = final_check(out_text)
    pages = count_pages()
    if pages is None:
        lines = out_text.count("\n") + 1
        pages = max(1, (lines + 49) // 50)

    print("OUTPUT", DST)
    print("console.log removed", STATS["console_log"])
    print("console.warn removed", STATS["console_warn"])
    print("console.error removed", STATS["console_error"])
    print("print removed", STATS["print"])
    print("audit/debug calls removed", STATS["audit_calls"])
    print("pages", pages)
    print("frontend_first", len(frontend) > 0 and len(backend) > 0)
    print("final_check", "PASS" if not issues else f"FAIL: {issues}")
    return 0 if not issues else 2


if __name__ == "__main__":
    sys.exit(main())
