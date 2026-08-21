#!/usr/bin/env python3
"""Purify gridvision_core_submit code for copyright filing. Does not touch main project."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
SRC_BACKEND = ROOT / "backend"
SRC_FRONTEND = ROOT / "frontend"
CLEAN = ROOT / "clean"
DOCX_OUT = ROOT / "GridVision_软著提交版_Clean.docx"
LINES_PER_PAGE = 50

STATS = {
    "console_log": 0,
    "console_warn": 0,
    "console_error": 0,
    "print": 0,
    "audit_functions_removed": 0,
    "audit_calls_removed": 0,
    "debug_comments_removed": 0,
    "ip_debug_removed": 0,
}

AUDIT_JS_FUNCTIONS = (
    "logInspectionPointCoordAudit",
    "logInspectionRouteBindingSummary",
    "formatInspectionCoordPair",
    "logRoutePointAudit",
    "printPlaybackCheck",
)

AUDIT_PY_CALL_PATTERN = re.compile(
    r"^\s*(_log_[a-zA-Z0-9_]+\([^)]*\)|_point_flow_log_dashboard\([^)]*\))\s*$"
)

DEBUG_COMMENT_PATTERN = re.compile(
    r"(DEBUG|AUDIT|TEMP|FIXME|TODO|专项排查|临时修复|兼容修复)",
    re.IGNORECASE,
)

IP_DEBUG_PATTERN = re.compile(r'IP_0005|IP_0007|IP_0012')


def count_matches(pattern: str, text: str) -> int:
    return len(re.findall(pattern, text))


def remove_js_function(text: str, name: str) -> str:
    global STATS
    pattern = rf"function\s+{re.escape(name)}\s*\("
    m = re.search(pattern, text)
    if not m:
        return text
    start = m.start()
    i = text.find("{", m.end())
    if i < 0:
        return text
    depth = 0
    j = i
    while j < len(text):
        c = text[j]
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                end = j + 1
                while end < len(text) and text[end] in " \t\r\n;":
                    if text[end] == "\n":
                        end += 1
                        break
                    end += 1
                STATS["audit_functions_removed"] += 1
                return text[:start] + text[end:]
        j += 1
    return text


def remove_console_calls(text: str) -> str:
    global STATS
    out = []
    i = 0
    n = len(text)
    while i < n:
        m = re.match(r"console\.(log|warn|error)\s*\(", text[i:])
        if not m:
            out.append(text[i])
            i += 1
            continue
        kind = m.group(1)
        STATS[f"console_{kind}"] += 1
        j = i + m.end()
        depth = 1
        in_str = None
        escape = False
        while j < n and depth > 0:
            c = text[j]
            if in_str:
                if escape:
                    escape = False
                elif c == "\\":
                    escape = True
                elif c == in_str:
                    in_str = None
            elif c in ("'", '"', "`"):
                in_str = c
            elif c == "(":
                depth += 1
            elif c == ")":
                depth -= 1
            j += 1
        while j < n and text[j] in " \t":
            j += 1
        if j < n and text[j] == ";":
            j += 1
        if j < n and text[j] == "\n":
            j += 1
        i = j
    return "".join(out)


def post_fix_js(text: str) -> str:
    text = re.sub(
        r"if \(dist <= BAD_SNAPPED_DIST_PX\) return true;[\s\S]*?return false;\n\}",
        "if (dist <= BAD_SNAPPED_DIST_PX) return true;\n  return false;\n}",
        text,
        count=1,
    )
    text = re.sub(
        r"\n  const first = sequence\[0\];\s*if \(first\) \{\s*\} else \{\s*\}\s*",
        "\n",
        text,
    )
    text = re.sub(
        r"getCurrentMission\(\);\s+const inspectionPointsLengthAtEntry",
        "getCurrentMission();\n  const inspectionPointsLengthAtEntry",
        text,
    )
    text = re.sub(r"\n\s{10,}failedPoints\.push", "\n      failedPoints.push", text)
    return text


def clean_js(text: str) -> str:
    global STATS
    for fn in AUDIT_JS_FUNCTIONS:
        text = remove_js_function(text, fn)

    # Remove logging-only blocks
    text = re.sub(
        r"\s*const _badSnappedLogged = new Set\(\);\s*\n",
        "\n",
        text,
        count=1,
    )
    text = re.sub(
        r"\s*_badSnappedLogged\.clear\(\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"  const pid = String\(p\.point_id \|\| p\.id \|\| \"\"\)\.trim\(\);\s*"
        r"if \(!_badSnappedLogged\.has\(pid\)\) \{[^}]+\}\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"    logRoutePointAudit\([^;]+\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"      logInspectionPointCoordAudit\(p\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"    logRoutePointAudit\([^;]+\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"\s*logInspectionRouteBindingSummary\([^;]+\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"\s*printPlaybackCheck\(\);\s*\n",
        "",
        text,
    )
    text = re.sub(
        r"\s*sequence\.forEach\(\(f, i\) => \{\s*"
        r"const ids = f\.routePointIds \|\| \[\];\s*"
        r"if \(ids\.length > 1\) \{[^}]+\}\s*\}\);\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"\s*expanded\.insertMeta\.forEach\(\(meta\) => \{[^}]+\}\);\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"\s*const segSummary = \[\];\s*let lastKey = null;\s*"
        r"sequence\.forEach\(\(f\) => \{[^}]+\}\);\s*"
        r"const first = sequence\[0\];\s*"
        r"if \(first\) \{[^}]+\} else \{[^}]+\}\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"\s*if \(finalCount\) \{\s*console\.log\([^}]+\}\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"\s*const tickPointIds = getFramePointIds\(ev\);\s*"
        r"if \(frameIndex % TICK_LOG_EVERY === 0 \|\| tickPointIds\.length\) \{[^}]+\}\s*",
        "",
        text,
        flags=re.DOTALL,
    )
    text = re.sub(
        r"  const redTraceFiniteCount = [^;]+;\s*console\.log\([^}]+\);\s*",
        "",
        text,
        flags=re.DOTALL,
    )

    if IP_DEBUG_PATTERN.search(text):
        STATS["ip_debug_removed"] += len(IP_DEBUG_PATTERN.findall(text))

    text = remove_console_calls(text)
    text = post_fix_js(text)

    lines = []
    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        if stripped.startswith("console.log(") or stripped.startswith("console.warn(") or stripped.startswith("console.error("):
            if "log" in stripped[:12]:
                STATS["console_log"] += stripped.startswith("console.log")
                STATS["console_warn"] += stripped.startswith("console.warn")
                STATS["console_error"] += stripped.startswith("console.error")
            continue
        if DEBUG_COMMENT_PATTERN.search(line) and (
            stripped.startswith("//") or stripped.startswith("*") or stripped.startswith("#")
        ):
            STATS["debug_comments_removed"] += 1
            continue
        lines.append(line)
    return "".join(lines)


def clean_py(text: str) -> str:
    global STATS
    lines_out = []
    skip_multiline_print = False
    paren_depth = 0

    for line in text.splitlines(keepends=True):
        stripped = line.strip()
        if skip_multiline_print:
            paren_depth += line.count("(") - line.count(")")
            if paren_depth <= 0:
                skip_multiline_print = False
            STATS["print"] += 1
            continue

        if AUDIT_PY_CALL_PATTERN.match(stripped):
            STATS["audit_calls_removed"] += 1
            continue

        if re.match(r"^\s*_log_[a-zA-Z0-9_]+\(", line):
            STATS["audit_calls_removed"] += 1
            if line.rstrip().endswith(")"):
                continue
            skip_multiline_print = True
            paren_depth = line.count("(") - line.count(")")
            continue

        if re.match(r"^\s*print\s*\(", line):
            STATS["print"] += 1
            if line.rstrip().endswith(")") and line.count("(") <= line.count(")"):
                continue
            skip_multiline_print = True
            paren_depth = line.count("(") - line.count(")")
            continue

        if DEBUG_COMMENT_PATTERN.search(line) and (
            stripped.startswith("#") or stripped.startswith('"""') or stripped.startswith("'''")
        ):
            STATS["debug_comments_removed"] += 1
            continue

        lines_out.append(line)

    return "".join(lines_out)


def estimate_pages(text: str) -> int:
    lines = max(1, text.count("\n"))
    return max(1, (lines + LINES_PER_PAGE - 1) // LINES_PER_PAGE)


def build_docx(files: list[tuple[str, str]], stats_report: str):
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING
    except ImportError:
        import subprocess
        subprocess.check_call(["pip", "install", "python-docx", "-q"])
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_LINE_SPACING

    doc = Document()
    title = doc.add_heading("GridVision 无人机巡检系统 — 软著提交核心源代码（Clean版）", 0)
    doc.add_paragraph(
        "本文档为 GridVision 系统核心源代码净化版，已移除调试日志、审计统计与专项排查代码，"
        "保留巡检点生成、任务规划、路径规划、重规划、Dashboard 导出、地图渲染、巡检播放与图传展示等核心业务逻辑。"
    )
    doc.add_heading("净化统计", level=1)
    doc.add_paragraph(stats_report)

    for label, content in files:
        doc.add_page_break()
        doc.add_heading(label, level=1)
        for para_text in content.split("\n"):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    doc.save(DOCX_OUT)


def main():
    if CLEAN.exists():
        shutil.rmtree(CLEAN)
    (CLEAN / "backend").mkdir(parents=True)
    (CLEAN / "frontend").mkdir(parents=True)

    before_text = ""
    after_text = ""
    docx_parts: list[tuple[str, str]] = []

    for src_dir, ext, cleaner in (
        (SRC_BACKEND, ".py", clean_py),
        (SRC_FRONTEND, ".js", clean_js),
    ):
        out_dir = CLEAN / src_dir.name
        for f in sorted(src_dir.glob(f"*{ext}")):
            raw = f.read_text(encoding="utf-8")
            before_text += raw
            cleaned = cleaner(raw)
            after_text += cleaned
            (out_dir / f.name).write_text(cleaned, encoding="utf-8")
            docx_parts.append((f"{src_dir.name}/{f.name}", cleaned))

    pages_before = estimate_pages(before_text)
    pages_after = estimate_pages(after_text)
    ratio = 100.0 * (1 - len(after_text) / max(1, len(before_text)))

    stats_md = f"""# GridVision 软著 Clean 版净化统计

## 删除项统计

| 类型 | 数量 |
|------|------|
| console.log | {STATS['console_log']} |
| console.warn | {STATS['console_warn']} |
| console.error | {STATS['console_error']} |
| print(...) | {STATS['print']} |
| audit 函数删除 | {STATS['audit_functions_removed']} |
| audit 调用删除 | {STATS['audit_calls_removed']} |
| debug 注释删除 | {STATS['debug_comments_removed']} |
| IP 专项排查痕迹 | {STATS['ip_debug_removed']} |

## 页数估算（按每页 {LINES_PER_PAGE} 行）

| 版本 | 估算页数 |
|------|----------|
| Clean 前 | {pages_before} |
| Clean 后 | {pages_after} |
| 精简比例 | {ratio:.1f}% |

## 输出

- 净化源码：`clean/backend/`, `clean/frontend/`
- Word 文档：`GridVision_软著提交版_Clean.docx`
"""
    (ROOT / "docs" / "clean_statistics.md").write_text(stats_md, encoding="utf-8")

    stats_report = (
        f"console.log: {STATS['console_log']}\n"
        f"console.warn: {STATS['console_warn']}\n"
        f"console.error: {STATS['console_error']}\n"
        f"print: {STATS['print']}\n"
        f"audit函数: {STATS['audit_functions_removed']}\n"
        f"audit调用: {STATS['audit_calls_removed']}\n"
        f"debug注释: {STATS['debug_comments_removed']}\n"
        f"Clean前页数(估): {pages_before}\n"
        f"Clean后页数(估): {pages_after}"
    )
    build_docx(docx_parts, stats_report)

    print(stats_md)
    print(f"Wrote {DOCX_OUT}")


if __name__ == "__main__":
    main()
